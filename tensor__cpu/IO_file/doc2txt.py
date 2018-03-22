# -*- coding: utf-8 -*-
from docx import Document

non_ansii_punctuation = {12290, 12289, 12300, 12301, 12302, 12303, 8216, 8217, 8220, 8221,
                         12308, 12309, 12304, 12305, 8212, 8230, 8211, 12298, 12299, 12296,
                         12297, 12288}

non_ansii_punctuation_dict = {'12290': 46, '12289': 44, '12300': 91, '12301': 93, '12302': 91, '12303': 93,
                              '8216': 39, '8217': 39, '8220': 34, '8221': 34, '12308': 91, '12309': 93,
                              '12304': 91, '12305': 93, '8212': 45, '8230': [46,46,46], '8211': 45, '12298': 34,
                              '12299': 34, '12296': 60, '12297': 62, '12288': 32}

def check_chinese(unicode_num):
    if 0x4E00 < unicode_num < 0x9FA5 or 0x9FA6 < unicode_num < 0x9FCB or 0x3400 < unicode_num < 0x4DB5 \
        or 0x20000 < unicode_num < 0x2A6D6 or 0x2A700 < unicode_num < 0x2B734 or 0x2B740 < unicode_num < 0x2B81D \
        or 0x2F00 < unicode_num < 0x2FD5 or 0x2E80 < unicode_num < 0x2EF3 or 0xF900 < unicode_num < 0xFAD9 \
        or 0x2F800 < unicode_num < 0x2FA1D or 0xE815 < unicode_num < 0xE86F or 0xE400 < unicode_num < 0xE5E8 \
        or 0xE600 < unicode_num < 0xE6CF or 0x31C0 < unicode_num < 0x31E3 or 0x2FF0 < unicode_num < 0x2FFB \
        or 0x3105 < unicode_num < 0x3120 or 0x31A0 < unicode_num < 0x31BA or unicode_num == 0x3007:

        return True

def strQ2B(ustring):
    """全角转半角"""
    rstring = ""
    for uchar in ustring:
        inside_code = ord(uchar)
        # print(uchar)
        # print(inside_code)
        if inside_code == 12288:  # 全角空格直接转换
            inside_code = 32
        elif inside_code == 8217:
            inside_code = 39 # '
        elif inside_code >= 65281 and inside_code <= 65374:  # 全角字符（除空格）根据关系转化
            inside_code -= 65248

        rstring += chr(inside_code)
    return rstring

def strQ2B_(ustring):
    """全角转半角"""
    rstring = ""
    for uchar in ustring:
        inside_code = ord(uchar)
        # print(uchar)
        # print(inside_code)
        if 65281 <= inside_code <= 65374:  # 全角字符（除空格）根据关系转化
            inside_code -= 65248
        elif str(inside_code) in non_ansii_punctuation_dict.keys():
            inside_code = non_ansii_punctuation_dict[str(inside_code)]

        rstring += chr(inside_code)
    return rstring




def strB2Q(ustring):
    """半角转全角"""
    rstring = ""
    for uchar in ustring:
        inside_code=ord(uchar)
        if inside_code == 32:                                 #半角空格直接转化
            inside_code = 12288
        elif inside_code >= 32 and inside_code <= 126:        #半角字符（除空格）根据关系转化
            inside_code += 65248

        rstring += chr(inside_code)
    return rstring




# document = Document('./xulian/24.0/16694168510253.docx')
with open('./xunlian_txt/0.0_32552700990073.txt', 'r') as f:
    with open('./txt_from_doc.txt', 'w') as fw:
        [fw.write(strQ2B_(f.read()))]

# with open('./txt_from_doc.txt', 'r', encoding='utf-8') as f:
#     for i in f.read():
#         print(ord(i))
#         print(i)

def doc2txt(doc_file, txt_file):
    print(doc_file)
    document = Document(doc_file)
    with open(txt_file, 'w', encoding='utf-8') as f_:
        [f_.write(strQ2B(paragraph.text + '\n')) for paragraph in document.paragraphs]